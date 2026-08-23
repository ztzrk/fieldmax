"""
FieldMax - Unhas Thesis Proposal Markdown to DOCX Converter
Strictly complies with:
PEDOMAN PENULISAN TUGAS AKHIR MAHASISWA UNIVERSITAS HASANUDDIN 2023
(SK Rektor No: 10438/UN4.1/KEP/2023)

Features:
- Precise Paragraph Alignment & Spacing Engine:
  * First paragraph after any heading: 0 cm indentation
  * Subsequent body paragraphs: 1.25 cm first-line indentation
  * Line spacing: 1.15 for body paragraphs, 1.0 for tables, captions, TOC, and math
- Advanced List & Numbering Hierarchy Engine:
  * Level 1 lists (1., 2., 3.): left_indent = 1.25 cm, first_line_indent = -0.6 cm
  * Level 2 nested lists (-, *, a., b.): left_indent = 2.0 cm, first_line_indent = -0.5 cm
  * Clean typographic bullet conversion (- / * -> •)
- Mathematical Equation Centering:
  * LaTeX math blocks ($$...$$) rendered as centered formulas with proper spacing
- Actual Visual Shapes inside UML and ERD Tables
- Strict Table Caption parsing (preventing paragraph descriptions from becoming captions)
- Advanced Table Layout Engine:
  * Strict proportional column widths matching 100% usable page width (B5 / A4)
  * Boxed academic borders (Top, Header Bottom, Bottom, Left, Right, Vertical Column Lines)
  * w:cantSplit on all rows to prevent mid-row page break splits
  * w:tblHeader on header rows for clean multi-page table continuation
- Native Microsoft Word Dynamic References:
  * DAFTAR ISI -> Word Table of Contents field (TOC \\o "1-3" \\h \\z \\u)
  * DAFTAR GAMBAR -> Word Table of Figures field (TOC \\h \\z \\c "Gambar")
  * DAFTAR TABEL -> Word Table of Figures field (TOC \\h \\z \\c "Tabel")
- Native Caption styles with Word SEQ fields (SEQ Gambar, SEQ Tabel)
- Native Microsoft Word Heading Styles (Heading 1, 2, 3, 4) for Navigation Pane
"""

import os
import re
import sys
import argparse
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from PIL import Image

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner cell padding in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_width(cell, width_cm):
    """Set explicit cell width in dxa."""
    width_dxa = int(width_cm * 567) # 1 cm = 567 dxa
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{width_dxa}" w:type="dxa"/>')
    tcPr.append(tcW)

def set_table_borders(table):
    """Apply boxed table grid border style with clear column and row lines."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="000000"/>\n'
        f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>\n'
        f'  <w:left w:val="single" w:sz="6" w:space="0" w:color="000000"/>\n'
        f'  <w:right w:val="single" w:sz="6" w:space="0" w:color="000000"/>\n'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="B0B0B0"/>\n'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="B0B0B0"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def format_run(run, bold=False, italic=False, font_name="Arial", font_size=10, color=RGBColor(0, 0, 0)):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    # Ensure EastAsian font is set as well
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.append(rFonts)

def clean_math_expression(math_str):
    """Convert LaTeX mathematical / logical expressions into clean Unicode notation."""
    math_clean = re.sub(r'\\text\{([^}]+)\}', r'\1', math_str)
    math_clean = math_clean.replace(r'\land', '∧').replace(r'\lor', '∨')
    math_clean = math_clean.replace(r'\rightarrow', '→').replace(r'\leftarrow', '←')
    math_clean = math_clean.replace(r'\le', '≤').replace(r'\ge', '≥')
    math_clean = math_clean.replace(r'\leq', '≤').replace(r'\geq', '≥')
    math_clean = math_clean.replace(r'\times', '×').replace(r'\neq', '≠')
    math_clean = math_clean.replace(r'\dots', '…').replace(r'\_', '_')
    return math_clean.strip()

def clean_markdown_text(text):
    """Remove HTML/SVG tags, anchors, and markdown links from text."""
    text = re.sub(r'<svg.*?</svg>', '', text, flags=re.DOTALL)
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
    # Only remove actual HTML tags, do not remove mathematical comparisons like < 10.00
    text = re.sub(r'</?[a-zA-Z][^>]*>', '', text)
    # Remove Obsidian block anchors like ^bab-1, ^gambar-10, ^tabel-4
    text = re.sub(r'\s*\^[a-zA-Z0-9_-]+', '', text)
    # Remove markdown link syntax [text](url) -> text
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    # Clean inline LaTeX math ($...$)
    text = re.sub(r'\$([^$\n]+)\$', lambda m: clean_math_expression(m.group(1)), text)
    return text.strip()

def resolve_image_path(raw_path, base_dir):
    """Resolve image path and map SVGs or drawio to converted PNGs."""
    raw_path = raw_path.strip().replace('\\', '/')
    possible_paths = [
        os.path.join(base_dir, raw_path),
        os.path.join(base_dir, 'images', os.path.basename(raw_path)),
        os.path.join(base_dir, 'images', 'symbols', os.path.basename(raw_path)),
        os.path.join(base_dir, os.path.basename(raw_path))
    ]
    
    # Check if direct file exists
    for p in possible_paths:
        if os.path.exists(p) and not p.endswith('.drawio') and not p.endswith('.svg'):
            return p
            
    # Check SVG replacement
    if raw_path.endswith('.svg'):
        png_path = raw_path[:-4] + '.png'
        for p in [os.path.join(base_dir, png_path), os.path.join(base_dir, 'images', os.path.basename(png_path))]:
            if os.path.exists(p):
                return p

    # Check drawio replacement to converted PNG
    if raw_path.endswith('.drawio'):
        png_path = raw_path[:-7] + '.png'
        for p in [os.path.join(base_dir, png_path), os.path.join(base_dir, 'images', os.path.basename(png_path)), os.path.join(base_dir, 'images', 'drawio', os.path.basename(png_path))]:
            if os.path.exists(p):
                return p
                
    # Check drawio replacements
    if 'erd' in raw_path:
        p = os.path.join(base_dir, 'images', 'image026.png')
        if os.path.exists(p): return p
    if 'relasi' in raw_path:
        p = os.path.join(base_dir, 'images', 'gambar-relasi-tabel.png')
        if os.path.exists(p): return p
    if 'booking-flow' in raw_path:
        p = os.path.join(base_dir, 'images', 'image019.png')
        if os.path.exists(p): return p
    if 'use-case' in raw_path:
        p = os.path.join(base_dir, 'images', 'image020.png')
        if os.path.exists(p): return p
    if 'alur-penelitian' in raw_path:
        p = os.path.join(base_dir, 'images', 'gambar-4-alur-penelitian.png')
        if os.path.exists(p): return p
    if 'waterfall' in raw_path:
        p = os.path.join(base_dir, 'images', 'gambar-waterfall-sdlc.png')
        if os.path.exists(p): return p
    if 'dsr' in raw_path:
        p = os.path.join(base_dir, 'images', 'gambar-dsr-fieldmax.png')
        if os.path.exists(p): return p
    if 'analisis-masalah' in raw_path:
        p = os.path.join(base_dir, 'images', 'gambar-analisis-masalah.png')
        if os.path.exists(p): return p
        
    for p in possible_paths:
        if os.path.exists(p):
            return p
    return None

def parse_markdown_spans(text, default_bold=False, default_italic=False):
    """
    Accurately parses inline markdown styles:
    - Code: `code`
    - Bold + Italic: ***bold italic*** or ___bold italic___
    - Bold with nested italic: **bold *italic* bold**
    - Bold: **bold** or __bold__
    - Italic: *italic* or _italic_
    Returns list of (text_chunk, is_bold, is_italic, is_code)
    """
    token_re = re.compile(
        r'(`[^`\n]+`)'                                  # 1: inline code
        r'|(\*\*\*[^*\n]+?\*\*\*)'                      # 2: triple asterisk bold-italic
        r'|(\*\*(?:[^*\n]|\*[^*\n]+?\*)+?\*\*)'         # 3: double asterisk bold (possibly containing *italic*)
        r'|(\*[^*\n]+?\*)'                              # 4: single asterisk italic
        r'|(_{3}[^_\n]+?_{3})'                          # 5: triple underscore bold-italic
        r'|(__[^_\n]+?__)'                              # 6: double underscore bold
        r'|(_[^_\n]+?_)'                                # 7: single underscore italic
    )

    tokens = []
    pos = 0

    for match in token_re.finditer(text):
        start, end = match.span()
        if start > pos:
            plain = text[pos:start]
            tokens.append((plain, default_bold, default_italic, False))

        full_match = match.group(0)

        # 1. Inline code
        if full_match.startswith('`') and full_match.endswith('`'):
            tokens.append((full_match[1:-1], default_bold, default_italic, True))

        # 2. Triple asterisk bold-italic
        elif full_match.startswith('***') and full_match.endswith('***'):
            tokens.append((full_match[3:-3], True, True, False))

        # 3. Triple underscore bold-italic
        elif full_match.startswith('___') and full_match.endswith('___'):
            tokens.append((full_match[3:-3], True, True, False))

        # 4. Double asterisk bold (with possible nested *italic*)
        elif full_match.startswith('**') and full_match.endswith('**'):
            inner = full_match[2:-2]
            # Check for nested italic inside bold
            sub_matches = list(re.finditer(r'(\*[^*\n]+?\*)', inner))
            if sub_matches:
                sub_pos = 0
                for sm in sub_matches:
                    s_start, s_end = sm.span()
                    if s_start > sub_pos:
                        tokens.append((inner[sub_pos:s_start], True, default_italic, False))
                    it_content = sm.group(1)[1:-1]
                    tokens.append((it_content, True, True, False))
                    sub_pos = s_end
                if sub_pos < len(inner):
                    tokens.append((inner[sub_pos:], True, default_italic, False))
            else:
                tokens.append((inner, True, default_italic, False))

        # 5. Double underscore bold
        elif full_match.startswith('__') and full_match.endswith('__'):
            tokens.append((full_match[2:-2], True, default_italic, False))

        # 6. Single asterisk italic
        elif full_match.startswith('*') and full_match.endswith('*'):
            tokens.append((full_match[1:-1], default_bold, True, False))

        # 7. Single underscore italic
        elif full_match.startswith('_') and full_match.endswith('_'):
            tokens.append((full_match[1:-1], default_bold, True, False))

        pos = end

    if pos < len(text):
        tokens.append((text[pos:], default_bold, default_italic, False))

    # Final cleanup: strip any stray unparsed asterisk/underscore characters in non-code chunks
    cleaned = []
    for chunk, b, it, c in tokens:
        if not c and ('*' in chunk or '_' in chunk):
            chunk_clean = re.sub(r'^\*+|\*+$', '', chunk)
            chunk_clean = re.sub(r'^_+|_+$', '', chunk_clean)
            if chunk_clean:
                cleaned.append((chunk_clean, b, it, c))
        else:
            if chunk:
                cleaned.append((chunk, b, it, c))

    return cleaned

def add_formatted_text(paragraph, text, base_bold=False, base_italic=False, font_size=10, font_name="Arial", base_dir=None):
    """Parse inline markdown (bold, italic, code, images) and add runs to paragraph."""
    # Check if text contains an inline image (e.g. inside table cell)
    img_match = re.search(r'!\[(.*?)\]\((.*?)\)|!\[\[(.*?)\]\]', text)
    if img_match and base_dir:
        raw_img = img_match.group(2) if img_match.group(2) else img_match.group(3)
        img_file = resolve_image_path(raw_img, base_dir)
        if img_file and os.path.exists(img_file):
            try:
                # Add picture into cell/paragraph
                paragraph.add_run().add_picture(img_file, width=Cm(1.2))
                return
            except Exception as e:
                print(f"Warning: cell picture {img_file}: {e}")

    text = clean_markdown_text(text)
    spans = parse_markdown_spans(text, default_bold=base_bold, default_italic=base_italic)
    for chunk, b, it, c in spans:
        run = paragraph.add_run(chunk)
        # Always use Arial across all runs (c code chunks use Arial italic)
        format_run(run, bold=b, italic=it or c, font_size=font_size, font_name="Arial")

def add_field_code(paragraph, instr_text):
    """Add a Microsoft Word dynamic field code (like TOC, SEQ, PAGEREF)."""
    r1 = paragraph.add_run()
    format_run(r1, font_name="Arial", font_size=10)
    fld1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    r1._r.append(fld1)

    r2 = paragraph.add_run()
    format_run(r2, font_name="Arial", font_size=10)
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> {instr_text} </w:instrText>')
    r2._r.append(instr)

    r3 = paragraph.add_run()
    format_run(r3, font_name="Arial", font_size=10)
    fld2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    r3._r.append(fld2)

    r4 = paragraph.add_run()
    format_run(r4, font_name="Arial", font_size=10)
    fld3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    r4._r.append(fld3)

def add_caption_with_seq(paragraph, caption_type, number_str, title_str, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """Add a native Word Caption paragraph containing SEQ field (SEQ Gambar / SEQ Tabel)."""
    paragraph.alignment = alignment
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(8) if caption_type == "Tabel" else Pt(2)
    paragraph.paragraph_format.space_after = Pt(2) if caption_type == "Tabel" else Pt(10)

    # Label: "Gambar " or "Tabel "
    run_lbl = paragraph.add_run(f"{caption_type} ")
    format_run(run_lbl, bold=True, font_size=10)

    # SEQ field
    r_seq1 = paragraph.add_run()
    format_run(r_seq1, font_name="Arial", font_size=10)
    fld_b = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    r_seq1._r.append(fld_b)

    r_seq2 = paragraph.add_run()
    format_run(r_seq2, font_name="Arial", font_size=10)
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> SEQ {caption_type} \\* ARABIC </w:instrText>')
    r_seq2._r.append(instr)

    r_seq3 = paragraph.add_run()
    format_run(r_seq3, font_name="Arial", font_size=10)
    fld_s = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    r_seq3._r.append(fld_s)

    # Fallback display number
    clean_num = re.sub(r'[^0-9a-zA-Z]', '', number_str)
    r_seq_disp = paragraph.add_run(clean_num)
    format_run(r_seq_disp, bold=True, font_size=10)

    r_seq4 = paragraph.add_run()
    format_run(r_seq4, font_name="Arial", font_size=10)
    fld_e = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    r_seq4._r.append(fld_e)

    # Dot separator
    run_dot = paragraph.add_run(". ")
    format_run(run_dot, bold=True, font_size=10)

    # Title text
    add_formatted_text(paragraph, title_str, base_bold=False, font_size=10)

def compute_column_widths(parsed_rows, usable_width_cm):
    """Compute optimal proportional column widths based on table structure and content."""
    num_cols = max(len(r) for r in parsed_rows)
    header_texts = [clean_markdown_text(c).upper() for c in parsed_rows[0]]

    # 2-column tables (e.g. Enum, Aktor)
    if num_cols == 2:
        return [usable_width_cm * 0.32, usable_width_cm * 0.68]

    # 3-column tables (e.g. Simbol UML, Verification Tokens, Sport Types, Sessions, Aktor)
    elif num_cols == 3:
        if "SIMBOL" in header_texts or "BENTUK" in header_texts[0]:
            return [usable_width_cm * 0.20, usable_width_cm * 0.25, usable_width_cm * 0.55]
        elif "NO" in header_texts[0]:
            return [usable_width_cm * 0.10, usable_width_cm * 0.30, usable_width_cm * 0.60]
        else:
            return [usable_width_cm * 0.26, usable_width_cm * 0.26, usable_width_cm * 0.48]

    # 4-column tables (Kamus Data OR Black Box Testing)
    elif num_cols == 4:
        if "NO" in header_texts[0] or "PENGUJIAN" in header_texts[1] or "DESKRIPSI" in header_texts[1]:
            # Black Box Testing Table
            return [usable_width_cm * 0.08, usable_width_cm * 0.38, usable_width_cm * 0.40, usable_width_cm * 0.14]
        else:
            # Database Schema Dictionary Table (Nama Field, Tipe Field, Keterangan, Default)
            return [usable_width_cm * 0.26, usable_width_cm * 0.22, usable_width_cm * 0.36, usable_width_cm * 0.16]

    # 5-column tables (e.g. Double Booking Prevention Test Matrix)
    elif num_cols == 5:
        return [usable_width_cm * 0.06, usable_width_cm * 0.24, usable_width_cm * 0.32, usable_width_cm * 0.26, usable_width_cm * 0.12]

    # 7-column table (Jadwal Pelaksanaan Penelitian)
    elif num_cols == 7:
        month_w = (usable_width_cm * 0.48) / 5.0
        return [usable_width_cm * 0.08, usable_width_cm * 0.44] + [month_w] * 5

    # General fallback
    else:
        col_w = usable_width_cm / float(num_cols)
        return [col_w] * num_cols

def setup_document_styles(doc):
    """Configure Word Document built-in styles according to Unhas 2023 guidelines."""
    # Enable automatic update of fields (TOC, Table of Figures) on open
    try:
        settings = doc.settings.element
        update_fields = parse_xml(f'<w:updateFields {nsdecls("w")} w:val="true"/>')
        settings.append(update_fields)
    except Exception as e:
        print(f"Notice: updateFields setting: {e}")

    # Set document-level default font in docDefaults to Arial
    try:
        styles_el = doc.styles.element
        docDefaults = styles_el.find(qn('w:docDefaults'))
        if docDefaults is None:
            docDefaults = parse_xml(f'<w:docDefaults {nsdecls("w")}/>')
            styles_el.insert(0, docDefaults)
        rPrDefault = docDefaults.find(qn('w:rPrDefault'))
        if rPrDefault is None:
            rPrDefault = parse_xml(f'<w:rPrDefault {nsdecls("w")}/>')
            docDefaults.append(rPrDefault)
        rPr = rPrDefault.find(qn('w:rPr'))
        if rPr is None:
            rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
            rPrDefault.append(rPr)
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial" w:eastAsia="Arial"/>')
        rPr.append(rFonts)
    except Exception as e:
        print(f"Notice: docDefaults font: {e}")

    # Normal Style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(0, 0, 0)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(0)
    normal_style.paragraph_format.space_before = Pt(0)

    # Heading 1 Style (Bab & Front Matter Main Headings)
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Arial'
    h1.font.size = Pt(11)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.line_spacing = 1.0
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.keep_with_next = True

    # Heading 2 Style (Sub-bab: 1.1, 1.2, 2.1, dst.)
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Arial'
    h2.font.size = Pt(10)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.line_spacing = 1.15
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.keep_with_next = True

    # Heading 3 Style (Anak Sub-bab: 1.6.1, 2.4.1, dst.)
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Arial'
    h3.font.size = Pt(10)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.line_spacing = 1.15
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(2)
    h3.paragraph_format.keep_with_next = True

    # Heading 4 Style (Sub-anak Sub-bab: #### 1. Next.js, dst.)
    h4 = doc.styles['Heading 4']
    h4.font.name = 'Arial'
    h4.font.size = Pt(10)
    h4.font.bold = True
    h4.font.color.rgb = RGBColor(0, 0, 0)
    h4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h4.paragraph_format.line_spacing = 1.15
    h4.paragraph_format.space_before = Pt(6)
    h4.paragraph_format.space_after = Pt(2)
    h4.paragraph_format.keep_with_next = True

    # Caption Style
    try:
        caption_style = doc.styles['Caption']
        caption_style.font.name = 'Arial'
        caption_style.font.size = Pt(10)
        caption_style.font.italic = False
        caption_style.font.color.rgb = RGBColor(0, 0, 0)
        caption_style.paragraph_format.line_spacing = 1.0
    except Exception:
        pass

def convert_md_to_docx(input_md_path, output_docx_path, page_size="B5"):
    print(f"[*] Reading source Markdown: {input_md_path}")
    with open(input_md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    base_dir = os.path.dirname(os.path.abspath(input_md_path))
    doc = Document()

    # Configure Section & Margins according to Unhas 2023 Guideline
    section = doc.sections[0]
    if page_size.upper() == "B5":
        section.page_width = Cm(17.6)
        section.page_height = Cm(25.0)
        usable_width_cm = 13.1 # 17.6 - 2.25 - 2.25
        max_img_width = Cm(13.0)
    else: # A4
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        usable_width_cm = 16.5 # 21.0 - 2.25 - 2.25
        max_img_width = Cm(15.5)

    # 5.1.2 Batas Sembir: 2.25 cm all sides
    section.top_margin = Cm(2.25)
    section.bottom_margin = Cm(2.25)
    section.left_margin = Cm(2.25)
    section.right_margin = Cm(2.25)

    # Setup styles
    setup_document_styles(doc)

    lines = content.split('\n')
    i = 0
    table_rows = []
    is_front_matter = True
    is_cover_page = True
    first_p_after_heading = False
    current_special_section = None # 'DAFTAR ISI', 'DAFTAR GAMBAR', 'DAFTAR TABEL'

    while i < len(lines):
        raw_line = lines[i]
        line = raw_line.strip()

        # Handle Markdown Tables
        if line.startswith('|') and line.endswith('|'):
            table_rows.append(line)
            i += 1
            continue
        elif table_rows:
            # If we are in DAFTAR ISI, DAFTAR GAMBAR, or DAFTAR TABEL, skip the static markdown table
            # and insert Word's dynamic Reference Field instead!
            if current_special_section == 'DAFTAR ISI':
                p_fld = doc.add_paragraph()
                p_fld.paragraph_format.line_spacing = 1.0
                add_field_code(p_fld, 'TOC \\o "1-3" \\h \\z \\u')
                current_special_section = None
                table_rows = []
                continue
            elif current_special_section == 'DAFTAR GAMBAR':
                p_fld = doc.add_paragraph()
                p_fld.paragraph_format.line_spacing = 1.0
                add_field_code(p_fld, 'TOC \\h \\z \\c "Gambar"')
                current_special_section = None
                table_rows = []
                continue
            elif current_special_section == 'DAFTAR TABEL':
                p_fld = doc.add_paragraph()
                p_fld.paragraph_format.line_spacing = 1.0
                add_field_code(p_fld, 'TOC \\h \\z \\c "Tabel"')
                current_special_section = None
                table_rows = []
                continue

            # Process normal tables
            parsed_rows = []
            for tr in table_rows:
                # Skip markdown separator row like |:---|:---| or | --- | --- |
                if re.match(r'^\|[\s\-:|]+\|$', tr):
                    continue
                cells = [c.strip() for c in tr.strip('|').split('|')]
                parsed_rows.append(cells)
            
            if parsed_rows:
                num_cols = max(len(r) for r in parsed_rows)
                col_widths = compute_column_widths(parsed_rows, usable_width_cm)
                
                tbl = doc.add_table(rows=len(parsed_rows), cols=num_cols)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_borders(tbl)

                for r_idx, row_data in enumerate(parsed_rows):
                    row = tbl.rows[r_idx]
                    
                    # cantSplit to prevent mid-row page splits
                    trPr = row._tr.get_or_add_trPr()
                    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
                    
                    # tblHeader on first row
                    if r_idx == 0:
                        trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

                    for c_idx in range(num_cols):
                        cell = row.cells[c_idx]
                        cell_text = row_data[c_idx] if c_idx < len(row_data) else ""
                        
                        # Set explicit width
                        w_cm = col_widths[c_idx] if c_idx < len(col_widths) else (usable_width_cm / num_cols)
                        set_cell_width(cell, w_cm)
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
                        
                        p = cell.paragraphs[0]
                        p.paragraph_format.line_spacing = 1.0
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                        
                        is_header = (r_idx == 0)
                        
                        # Alignment rules:
                        # Headers: Center
                        # Numbers / Booleans / Results / Images: Center
                        # Descriptions / Names: Left
                        if is_header:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            tcPr = cell._tc.get_or_add_tcPr()
                            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
                            tcPr.append(shd)
                        else:
                            if '![' in cell_text or '![[' in cell_text:
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            elif c_idx == 0 and num_cols >= 4 and cell_text.isdigit():
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            elif c_idx == num_cols - 1 and cell_text.upper() in ["BERHASIL", "GAGAL", "TRUE", "FALSE", "USER", "RENTER", "ADMIN", "NOW()", "UUID()", "NULLABLE"]:
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            elif num_cols == 7 and c_idx >= 2: # Schedule checkmarks
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            else:
                                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                
                        add_formatted_text(p, cell_text, base_bold=is_header, font_size=9 if num_cols > 3 else 9.5, base_dir=base_dir)

                # Add a small spacing after table
                sp_p = doc.add_paragraph()
                sp_p.paragraph_format.space_before = Pt(0)
                sp_p.paragraph_format.space_after = Pt(4)

            table_rows = []

        if not line:
            i += 1
            continue

        # Skip horizontal rules
        if line == '---' or line == '***':
            i += 1
            continue

        # Check for LaTeX Math Formula Blocks ($$ ... $$)
        if line.startswith('$$') and line.endswith('$$') and len(line) > 4:
            math_content = line[2:-2].strip()
            math_display = clean_math_expression(math_content)
            
            p_math = doc.add_paragraph()
            p_math.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_math.paragraph_format.line_spacing = 1.15
            p_math.paragraph_format.space_before = Pt(6)
            p_math.paragraph_format.space_after = Pt(6)
            p_math.paragraph_format.first_line_indent = Cm(0)
            
            run_m = p_math.add_run(math_display)
            format_run(run_m, bold=True, italic=True, font_name="Arial", font_size=10)
            i += 1
            continue

        # Check for Chapter Title (# BAB I / # BAB II / # ABSTRAK / # DAFTAR ISI etc.)
        if line.startswith('# '):
            heading_text = line[2:].strip()
            heading_clean = clean_markdown_text(heading_text)
            
            # Handle Cover Title differently
            if is_cover_page and ("RANCANG BANGUN" in heading_clean.upper()):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(12)
                run = p.add_run(heading_clean.upper())
                format_run(run, bold=True, font_size=11)
                i += 1
                continue

            # Formal Headings (Heading 1)
            if "BAB " in heading_clean.upper() or any(k in heading_clean.upper() for k in ["ABSTRAK", "ABSTRACT", "DAFTAR ISI", "DAFTAR GAMBAR", "DAFTAR TABEL", "DAFTAR PUSTAKA"]):
                is_cover_page = False
                if i > 5: # Page break before each main chapter
                    doc.add_page_break()

                # Mark special sections for Reference Fields
                if "DAFTAR ISI" in heading_clean.upper():
                    current_special_section = 'DAFTAR ISI'
                elif "DAFTAR GAMBAR" in heading_clean.upper():
                    current_special_section = 'DAFTAR GAMBAR'
                elif "DAFTAR TABEL" in heading_clean.upper():
                    current_special_section = 'DAFTAR TABEL'
                else:
                    current_special_section = None

                # Add as native Heading 1 so it appears in Word Navigation Pane
                p = doc.add_paragraph(style='Heading 1')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(12)
                
                # Split BAB I PENDAHULUAN into two lines for display while keeping it in Navigation pane
                bab_match = re.match(r'^(BAB\s+[IVXLCDM]+)\s+(.*)$', heading_clean, re.IGNORECASE)
                if bab_match:
                    run1 = p.add_run(bab_match.group(1).upper() + "\n")
                    format_run(run1, bold=True, font_size=11)
                    run2 = p.add_run(bab_match.group(2).upper())
                    format_run(run2, bold=True, font_size=11)
                    is_front_matter = False
                else:
                    run = p.add_run(heading_clean.upper())
                    format_run(run, bold=True, font_size=11)
                
                first_p_after_heading = True
                i += 1
                continue

        # Check for Section Headings (## 1.1, ## 2.1, etc.) -> Heading 2
        if line.startswith('## '):
            heading_text = line[3:].strip()
            heading_clean = clean_markdown_text(heading_text)
            p = doc.add_paragraph(style='Heading 2')
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            add_formatted_text(p, heading_clean, base_bold=True, font_size=10, base_dir=base_dir)
            first_p_after_heading = True
            i += 1
            continue

        # Check for Sub-section Headings (### 1.6.1, ### 2.4.1, etc.) -> Heading 3
        if line.startswith('### '):
            heading_text = line[4:].strip()
            heading_clean = clean_markdown_text(heading_text)
            
            # If on cover page (e.g. ### **PROPOSAL SKRIPSI**)
            if is_cover_page:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                add_formatted_text(p, heading_clean, base_bold=True, font_size=10, base_dir=base_dir)
            else:
                p = doc.add_paragraph(style='Heading 3')
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(2)
                add_formatted_text(p, heading_clean, base_bold=True, font_size=10, base_dir=base_dir)
            first_p_after_heading = True
            i += 1
            continue

        # Check for Sub-sub-section Headings (#### 1. Next.js, etc.) -> Heading 4
        if line.startswith('#### '):
            heading_text = line[5:].strip()
            heading_clean = clean_markdown_text(heading_text)
            p = doc.add_paragraph(style='Heading 4')
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            add_formatted_text(p, heading_clean, base_bold=True, font_size=10, base_dir=base_dir)
            first_p_after_heading = True
            i += 1
            continue

        # Check for Table Titles (e.g. **Tabel 4.** Judul Tabel) -> Native Caption with SEQ Tabel
        # Strict matching: requires dot inside bold asterisks e.g. **Tabel 13.** or **Tabel 4a.**
        tbl_caption_match = re.match(r'^\*\*(Tabel\s+([0-9a-zA-Z]+(?:\.[0-9]+)?)\.)\*\*\s*(.*)$', line)
        if tbl_caption_match:
            num_str = tbl_caption_match.group(2).strip()
            title = clean_markdown_text(tbl_caption_match.group(3).strip()).rstrip('.')
            
            # Ensure it is a genuine short title, not a body sentence starting with a reference
            if len(title) <= 120 and not title.startswith('('):
                p = doc.add_paragraph(style='Caption')
                add_caption_with_seq(p, caption_type="Tabel", number_str=num_str, title_str=title, alignment=WD_ALIGN_PARAGRAPH.LEFT)
                i += 1
                continue

        # Check for Images (![alt](path) or ![[path]])
        img_match = re.search(r'!\[(.*?)\]\((.*?)\)|!\[\[(.*?)\]\]', line)
        if img_match:
            raw_img_path = img_match.group(2) if img_match.group(2) else img_match.group(3)
            img_file = resolve_image_path(raw_img_path, base_dir)
            if img_file and os.path.exists(img_file):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(4)
                
                # Check dimensions and scale
                try:
                    with Image.open(img_file) as im:
                        w, h = im.size
                        r_pic = p.add_run()
                        format_run(r_pic, font_name="Arial", font_size=10)
                        # If cover logo, keep it neat around 3.5 cm
                        if 'image001' in img_file:
                            r_pic.add_picture(img_file, width=Cm(3.5))
                        elif 'activity' in img_file:
                            aspect = h / float(w)
                            target_w_cm = min(9.0, 14.5 / aspect)
                            r_pic.add_picture(img_file, width=Cm(target_w_cm))
                        else:
                            target_w = max_img_width
                            if w < 600:
                                target_w = Cm(min(max_img_width.cm, w / 37.79))
                            r_pic.add_picture(img_file, width=target_w)
                except Exception as e:
                    print(f"Warning: could not insert image {img_file}: {e}")
            i += 1
            continue

        # Check for Figure Captions (e.g. **Gambar 1.** Judul Gambar) -> Native Caption with SEQ Gambar
        # Strict matching: requires dot inside bold asterisks e.g. **Gambar 1.** or **Gambar 9a.**
        fig_caption_match = re.match(r'^\*\*(Gambar\s+([0-9a-zA-Z]+(?:\.[0-9]+)?)\.)\*\*\s*(.*)$', line)
        if fig_caption_match:
            num_str = fig_caption_match.group(2).strip()
            title = clean_markdown_text(fig_caption_match.group(3).strip()).rstrip('.')
            
            if len(title) <= 150:
                p = doc.add_paragraph(style='Caption')
                add_caption_with_seq(p, caption_type="Gambar", number_str=num_str, title_str=title, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                i += 1
                continue

        # Handle Standard Paragraphs
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Check if line is centered in HTML block or Title block
        if '<div align="center">' in line or '<div align=\'center\'>' in line:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        if '</div>' in line or '<br>' in line:
            i += 1
            continue

        if is_cover_page:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            add_formatted_text(p, line, font_size=10, base_dir=base_dir)
            i += 1
            continue

        # Advanced List Item Detection (Numbered, Bullets, Sub-items, Nested Lists)
        list_match = re.match(r'^(\s*)(\d+[\.\)]|[\*\-\•]|[a-zA-Z][\.\)])\s+(.*)$', raw_line)
        if list_match:
            indent_spaces = list_match.group(1)
            bullet = list_match.group(2)
            rest = list_match.group(3)
            is_nested = (len(indent_spaces) >= 2 or bullet in ['-', '*', '•'])

            if is_nested:
                p.paragraph_format.left_indent = Cm(1.85)
                p.paragraph_format.first_line_indent = Cm(-0.5)
            else:
                p.paragraph_format.left_indent = Cm(1.25)
                p.paragraph_format.first_line_indent = Cm(-0.6)

            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)

            # Format bullet indicator
            if bullet in ['-', '*']:
                run_b = p.add_run("• ")
                format_run(run_b, bold=False, font_size=10)
            else:
                run_b = p.add_run(f"{bullet} ")
                is_num = bool(re.match(r'^\d+\.?$', bullet))
                format_run(run_b, bold=is_num, font_size=10)

            add_formatted_text(p, rest, font_size=10, base_dir=base_dir)
        else:
            # Body paragraph: First paragraph after heading has 0 indent, others have 1.25 cm
            if not is_front_matter:
                p.paragraph_format.first_line_indent = Cm(0) if first_p_after_heading else Cm(1.25)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            add_formatted_text(p, line, font_size=10, base_dir=base_dir)

        first_p_after_heading = False
        i += 1

    print(f"[*] Saving DOCX document to: {output_docx_path}")
    doc.save(output_docx_path)
    print(f"[+] Conversion complete! Output size: {os.path.getsize(output_docx_path):,} bytes")

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description="Convert Markdown Skripsi to DOCX following Unhas 2023 Guidelines with Word Dynamic References and Fixed Table Layout")
    parser.add_argument("--input", default=r"c:\Users\Ztzrk\Documents\fieldmax\obsidian\Fieldmax\Skripsi Lengkap - BAB I s.d. BAB IV (Audited).md", help="Path to input Markdown file")
    parser.add_argument("--output", default=r"c:\Users\Ztzrk\Documents\fieldmax\obsidian\Fieldmax\Skripsi Lengkap - BAB I s.d. BAB IV (Audited).docx", help="Path to output DOCX file")
    parser.add_argument("--size", default="B5", choices=["B5", "A4"], help="Paper size (B5 as per Unhas guideline, or A4)")
    args = parser.parse_args()

    convert_md_to_docx(args.input, args.output, page_size=args.size)
